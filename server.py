from dotenv import load_dotenv
import os
import io
from flask import Flask, request, jsonify
from flask_cors import CORS
import google.generativeai as genai
import base64
import docx  # Import the python-docx library

dotenv_path = os.path.join(os.path.dirname(__file__), '.env')  # Path to .env file
load_dotenv(dotenv_path)

print("API Key:", os.environ.get("GOOGLE_API_KEY")) # Add this line

app = Flask(__name__)
CORS(app)

# Configure Gemini API
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("No Google API key found.  Make sure you have set the GOOGLE_API_KEY environment variable.")
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')

# Load knowledge base
knowledge_base_text = ""
knowledge_base_path = os.path.join(os.path.dirname(__file__), 'knowledge_base')
for filename in os.listdir(knowledge_base_path):
    if filename.endswith(".docx"):
        filepath = os.path.join(knowledge_base_path, filename)
        try:
            document = docx.Document(filepath)
            text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            knowledge_base_text += f"\n\n--- Content from {filename} ---\n{text}"
        except Exception as e:
            print(f"Error reading {filename}: {e}")

# Store conversations
conversations = {}

def process_image_data(data):
    parts = data.split(',')
    if len(parts) > 1:
        base64_data = parts[1]
        image_data = base64.b64decode(base64_data)
        return {"mime_type": "image/png", "data": image_data}
    return None

def extract_text_from_docx(docx_data):
    try:
        document = docx.Document(io.BytesIO(docx_data))
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None

@app.route('/run', methods=['POST'])
def run_model():
    data = request.get_json()
    prompt = data.get('task')
    conversation_id = data.get('conversationId')
    image_data = data.get('fileData')

    if not prompt:
        return jsonify({'error': 'Prompt is required'}), 400

    # Get or create conversation
    if conversation_id not in conversations:
        conversations[conversation_id] = model.start_chat(history=[])
    conversation = conversations[conversation_id]

    try:
        # System prompt
        system_prompt = """
        Hey there! I'm the TLO Chatbot, your friendly guide to all things Thanzi La Onse (TLO) Model.
        I'm here to help you understand the TLO model, which is part of the Thanzi La Onse (Health for All) Collaboration, funded by Wellcome (and formerly GCRF and UKRI).

        Feel free to ask me anything about the TLO model, its code base, or its applications.
        I'll do my best to provide you with detailed and accurate information, drawing from my knowledge base.
        If I don't know the answer, I'll let you know!
        """

        # Add knowledge base to prompt
        full_prompt = f"{system_prompt}\n\nKNOWLEDGE BASE:\n{knowledge_base_text}\n\nUSER QUESTION:\n{prompt}"

        content = [full_prompt]
        if image_data:
            image = process_image_data(image_data)
            if image:
                content.append(image)
            else:
                return jsonify({'error': 'Invalid image data'}), 400

        # Extract text from DOCX file
        if image_data and image_data.startswith("data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,"):
            docx_data = base64.b64decode(image_data.split(",")[1])
            file_text = extract_text_from_docx(docx_data)
            if file_text:
                full_prompt += f"\n\nCONTENT OF UPLOADED FILE:\n{file_text}"

        response = conversation.send_message(content)
        return jsonify({'final_result': response.text})
    except Exception as e:
        print(f"Error generating content: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)